"""
ATS Resume Scanner - Simplified FastAPI Backend
Production-grade API without heavy ML dependencies for immediate deployment
"""

import os
import re
import io
from typing import Optional, List, Dict, Any
from contextlib import asynccontextmanager
from dataclasses import dataclass

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import pdfplumber
from docx import Document

# ============== Pydantic Models ==============

class ScoreBreakdown(BaseModel):
    keyword_match: float
    semantic_match: float
    experience_alignment: float
    formatting_compatibility: float
    section_completeness: float

class MatchedKeyword(BaseModel):
    keyword: str
    matched: bool
    importance: str = "medium"

class FormattingIssue(BaseModel):
    type: str
    severity: str
    message: str
    suggestion: str = ""

class OptimizationSuggestion(BaseModel):
    category: str
    priority: str
    message: str
    action_items: List[str] = []

class ATSScanResponse(BaseModel):
    ats_score: float
    score_breakdown: ScoreBreakdown
    matched_keywords: List[MatchedKeyword]
    missing_keywords: List[str]
    formatting_issues: List[FormattingIssue]
    optimization_suggestions: List[OptimizationSuggestion]
    extracted_data: Dict[str, Any]
    resume_text_preview: str

class HealthResponse(BaseModel):
    status: str
    version: str

# ============== Resume Parser ==============

class ResumeParser:
    def parse(self, content: bytes, filename: str) -> Dict[str, Any]:
        ext = filename.lower().split('.')[-1]
        
        if ext == 'pdf':
            return self._parse_pdf(content)
        elif ext in ['docx', 'doc']:
            return self._parse_docx(content)
        elif ext == 'txt':
            return self._parse_txt(content)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def _parse_pdf(self, content: bytes) -> Dict[str, Any]:
        text_parts = []
        has_tables = False
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
                if page.extract_tables():
                    has_tables = True
        
        return {
            'text': '\n'.join(text_parts),
            'has_tables': has_tables,
            'pages': len(pdf.pages)
        }
    
    def _parse_docx(self, content: bytes) -> Dict[str, Any]:
        doc = Document(io.BytesIO(content))
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        return {
            'text': text,
            'has_tables': len(doc.tables) > 0,
            'pages': 1
        }
    
    def _parse_txt(self, content: bytes) -> Dict[str, Any]:
        text = content.decode('utf-8', errors='ignore')
        return {
            'text': text,
            'has_tables': False,
            'pages': 1
        }

# ============== Entity Extractor ==============

class EntityExtractor:
    SKILLS = [
        'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'go', 'rust',
        'react', 'angular', 'vue', 'nodejs', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
        'git', 'github', 'gitlab', 'jenkins', 'circleci',
        'html', 'css', 'sass', 'less', 'tailwind',
        'machine learning', 'data science', 'ai', 'nlp',
        'rest api', 'graphql', 'microservices', 'serverless'
    ]
    
    JOB_TITLES = [
        'software engineer', 'developer', 'programmer', 'coder',
        'full stack', 'frontend', 'backend', 'devops',
        'data scientist', 'data engineer', 'ml engineer',
        'product manager', 'project manager', 'tech lead',
        'qa engineer', 'test engineer', 'security engineer'
    ]
    
    CERTIFICATIONS = [
        'aws certified', 'azure certified', 'gcp certified',
        'pmp', 'cissp', 'ceh', 'security+', 'network+',
        'scrum master', 'csm', 'psm', 'itil',
        'ccna', 'ccnp', 'oracle certified'
    ]
    
    def extract(self, text: str) -> Dict[str, Any]:
        text_lower = text.lower()
        
        skills = [s for s in self.SKILLS if s in text_lower]
        titles = [t for t in self.JOB_TITLES if t in text_lower]
        certs = [c for c in self.CERTIFICATIONS if c in text_lower]
        
        # Extract years of experience
        years = 0
        matches = re.findall(r'(\d+)\+?\s*years?\s+(?:of\s+)?experience', text_lower)
        if matches:
            years = max(int(m) for m in matches)
        
        # Extract email
        email = ''
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            email = email_match.group(0)
        
        # Extract phone
        phone = ''
        phone_match = re.search(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text)
        if phone_match:
            phone = phone_match.group(0)
        
        return {
            'skills': skills,
            'job_titles': titles,
            'certifications': certs,
            'years_experience': years,
            'email': email,
            'phone': phone
        }

# ============== Keyword Matcher ==============

class KeywordMatcher:
    SYNONYMS = {
        'python': ['py'],
        'javascript': ['js', 'node.js', 'nodejs'],
        'react': ['reactjs', 'react.js'],
        'aws': ['amazon web services'],
        'docker': ['containerization'],
        'kubernetes': ['k8s'],
    }
    
    def match(self, resume_text: str, jd_keywords: List[str], resume_entities: Dict) -> Dict[str, Any]:
        resume_lower = resume_text.lower()
        resume_skills = set(resume_entities.get('skills', []))
        
        matches = []
        missing = []
        
        for keyword in jd_keywords:
            kw_lower = keyword.lower()
            matched = False
            match_type = 'none'
            
            # Exact match
            if kw_lower in resume_lower or kw_lower in resume_skills:
                matched = True
                match_type = 'exact'
            # Synonym match
            elif kw_lower in self.SYNONYMS:
                for syn in self.SYNONYMS[kw_lower]:
                    if syn in resume_lower:
                        matched = True
                        match_type = 'synonym'
                        break
            
            if matched:
                matches.append({'keyword': keyword, 'matched': True, 'type': match_type})
            else:
                missing.append(keyword)
        
        total = len(jd_keywords) if jd_keywords else 1
        score = (len(matches) / total) * 100
        
        return {
            'matches': matches,
            'missing': missing,
            'score': round(score, 1)
        }

# ============== Formatting Analyzer ==============

class FormattingAnalyzer:
    def analyze(self, parsed_data: Dict, text: str) -> Dict[str, Any]:
        issues = []
        
        # Check for tables
        if parsed_data.get('has_tables'):
            issues.append({
                'type': 'table',
                'severity': 'warning',
                'message': 'Tables detected in resume',
                'suggestion': 'Convert tables to plain text for better ATS compatibility'
            })
        
        # Check for images (heuristic)
        if '[image]' in text.lower() or 'img' in text.lower():
            issues.append({
                'type': 'image',
                'severity': 'warning',
                'message': 'Possible images detected',
                'suggestion': 'Remove images and use text only'
            })
        
        # Check for contact info
        if not re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text):
            issues.append({
                'type': 'contact',
                'severity': 'critical',
                'message': 'No email address found',
                'suggestion': 'Add a professional email address'
            })
        
        # Check sections
        sections_found = []
        text_lower = text.lower()
        if 'experience' in text_lower or 'work' in text_lower:
            sections_found.append('experience')
        if 'education' in text_lower:
            sections_found.append('education')
        if 'skills' in text_lower:
            sections_found.append('skills')
        
        # Score based on issues and sections
        score = 100 - len([i for i in issues if i['severity'] == 'critical']) * 15
        score -= len([i for i in issues if i['severity'] == 'warning']) * 5
        score += len(sections_found) * 5
        score = max(0, min(100, score))
        
        return {
            'score': round(score, 1),
            'issues': issues,
            'sections_found': sections_found
        }

# ============== ATS Scorer ==============

class ATSScorer:
    WEIGHTS = {
        'keyword': 0.40,
        'semantic': 0.25,
        'experience': 0.20,
        'format': 0.10,
        'section': 0.05
    }
    
    def calculate(self, keyword_score: float, format_score: float, 
                  sections_found: List[str], resume_entities: Dict) -> Dict[str, Any]:
        
        # Semantic score (simplified - based on content richness)
        text_length = len(resume_entities.get('skills', [])) + len(resume_entities.get('job_titles', []))
        semantic_score = min(100, text_length * 10)
        
        # Experience score
        years = resume_entities.get('years_experience', 0)
        exp_score = min(100, years * 15 + 30)
        
        # Section score
        required = ['experience', 'education']
        section_score = sum(20 for s in required if s in sections_found)
        section_score += 10 if 'skills' in sections_found else 0
        
        # Calculate weighted total
        total = (
            keyword_score * self.WEIGHTS['keyword'] +
            semantic_score * self.WEIGHTS['semantic'] +
            exp_score * self.WEIGHTS['experience'] +
            format_score * self.WEIGHTS['format'] +
            section_score * self.WEIGHTS['section']
        )
        
        # Generate suggestions
        suggestions = []
        
        if keyword_score < 70:
            suggestions.append({
                'category': 'keywords',
                'priority': 'high',
                'message': 'Improve keyword matching by adding relevant skills from the job description',
                'action_items': ['Review job description for required skills', 'Add missing technical skills']
            })
        
        if format_score < 80:
            suggestions.append({
                'category': 'format',
                'priority': 'medium',
                'message': 'Address formatting issues for better ATS compatibility',
                'action_items': ['Remove tables and complex formatting', 'Use standard fonts']
            })
        
        if 'experience' not in sections_found:
            suggestions.append({
                'category': 'structure',
                'priority': 'high',
                'message': 'Add an Experience section',
                'action_items': ['Include work history with dates', 'Add bullet points with achievements']
            })
        
        return {
            'total_score': round(total, 1),
            'keyword_score': round(keyword_score, 1),
            'semantic_score': round(semantic_score, 1),
            'experience_score': round(exp_score, 1),
            'format_score': round(format_score, 1),
            'section_score': round(section_score, 1),
            'suggestions': suggestions
        }

# ============== FastAPI App ==============

app = FastAPI(
    title="ATS Resume Scanner API",
    description="Production-grade ATS resume analysis",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize components
parser = ResumeParser()
extractor = EntityExtractor()
keyword_matcher = KeywordMatcher()
format_analyzer = FormattingAnalyzer()
scorer = ATSScorer()

@app.get("/api/v1/health", response_model=HealthResponse)
async def health():
    return HealthResponse(status="healthy", version="1.0.0")

@app.post("/api/v1/ats/scan", response_model=ATSScanResponse)
async def scan_resume(
    resume: UploadFile = File(...),
    job_description: Optional[str] = Form(None)
):
    # Validate file
    allowed = ['application/pdf', 'application/msword', 
               'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
               'text/plain']
    
    if resume.content_type not in allowed:
        raise HTTPException(400, "Invalid file type. Use PDF, DOCX, or TXT.")
    
    content = await resume.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(400, "File too large. Max 5MB.")
    
    # Parse resume
    parsed = parser.parse(content, resume.filename)
    resume_text = parsed['text']
    
    # Extract entities
    entities = extractor.extract(resume_text)
    
    # Extract JD keywords if provided
    jd_keywords = []
    if job_description:
        jd_entities = extractor.extract(job_description)
        jd_keywords = jd_entities['skills'] + jd_entities['job_titles']
    
    # Match keywords
    keyword_result = keyword_matcher.match(resume_text, jd_keywords, entities)
    
    # Analyze formatting
    format_result = format_analyzer.analyze(parsed, resume_text)
    
    # Calculate score
    score_result = scorer.calculate(
        keyword_result['score'],
        format_result['score'],
        format_result['sections_found'],
        entities
    )
    
    return ATSScanResponse(
        ats_score=score_result['total_score'],
        score_breakdown=ScoreBreakdown(
            keyword_match=score_result['keyword_score'],
            semantic_match=score_result['semantic_score'],
            experience_alignment=score_result['experience_score'],
            formatting_compatibility=score_result['format_score'],
            section_completeness=score_result['section_score']
        ),
        matched_keywords=[
            MatchedKeyword(keyword=m['keyword'], matched=True, importance='medium')
            for m in keyword_result['matches']
        ],
        missing_keywords=keyword_result['missing'],
        formatting_issues=[
            FormattingIssue(type=i['type'], severity=i['severity'], 
                          message=i['message'], suggestion=i.get('suggestion', ''))
            for i in format_result['issues']
        ],
        optimization_suggestions=[
            OptimizationSuggestion(
                category=s['category'],
                priority=s['priority'],
                message=s['message'],
                action_items=s.get('action_items', [])
            )
            for s in score_result['suggestions']
        ],
        extracted_data={
            'skills': entities['skills'],
            'job_titles': entities['job_titles'],
            'certifications': entities['certifications'],
            'years_experience': entities['years_experience']
        },
        resume_text_preview=resume_text[:500] + "..." if len(resume_text) > 500 else resume_text
    )

@app.post("/api/v1/ats/parse-resume")
async def parse_resume(resume: UploadFile = File(...)):
    content = await resume.read()
    parsed = parser.parse(content, resume.filename)
    entities = extractor.extract(parsed['text'])
    format_result = format_analyzer.analyze(parsed, parsed['text'])
    
    return {
        'raw_text': parsed['text'],
        'extracted_data': entities,
        'formatting_analysis': format_result,
        'sections_found': format_result['sections_found']
    }

@app.post("/api/v1/ats/extract-keywords")
async def extract_keywords(job_description: str = Form(...)):
    entities = extractor.extract(job_description)
    return {
        'skills': entities['skills'],
        'job_titles': entities['job_titles'],
        'certifications': entities['certifications'],
        'keywords': entities['skills'] + entities['job_titles']
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
